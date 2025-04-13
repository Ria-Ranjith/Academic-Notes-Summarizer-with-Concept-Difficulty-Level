<<<<<<< HEAD
from pdfminer.high_level import extract_text
import docx
from typing import Union, IO
import warnings

def extract_text_from_pdf(file: Union[str, IO[bytes]]) -> str:
    """
    Extracts text from a PDF file.
    
    Args:
        file: Either a file path (str) or a file-like object (BytesIO).
    
    Returns:
        Extracted text as a string.
    """
    try:
        with warnings.catch_warnings():
            warnings.simplefilter("ignore")  # Suppress PDFMiner warnings
            text = extract_text(file)
        return text.strip() if text else ""
    except Exception as e:
        raise RuntimeError(f"Failed to extract PDF text: {str(e)}")

def extract_text_from_docx(file: Union[str, IO[bytes]]) -> str:
    """
    Extracts text from a DOCX file.
    
    Args:
        file: Either a file path (str) or a file-like object (BytesIO).
    
    Returns:
        Extracted text as a string.
    """
    try:
        doc = docx.Document(file)
        return '\n'.join(para.text for para in doc.paragraphs if para.text.strip())
    except Exception as e:
        raise RuntimeError(f"Failed to extract DOCX text: {str(e)}")
=======
# utils.py

from pdfminer.high_level import extract_text
from docx import Document
from langchain.prompts import PromptTemplate
from langchain.llms import OpenAI
from langchain.chains import LLMChain

# Function to extract text from PDF
def extract_pdf_text(file_path):
    return extract_text(file_path)

# Function to extract text from DOCX
def extract_docx_text(file_path):
    doc = Document(file_path)
    text = ''
    for para in doc.paragraphs:
        text += para.text + '\n'
    return text

# Define the prompt template
prompt_template = """Given the following academic notes: {notes_content},
summarize the main points and create flashcards with questions and answers for study purposes."""

prompt = PromptTemplate(input_variables=["notes_content"], template=prompt_template)

# Set up the LLM (language model) using OpenAI's API
llm = OpenAI(temperature=0.7)

# Create the Langchain chain
chain = LLMChain(llm=llm, prompt=prompt)

# Function to process the notes and generate summary and flashcards
def generate_summary_and_flashcards(notes_content):
    # Run the chain with the notes content
    response = chain.run(notes_content)
    return response
>>>>>>> 51e8a6e42e67735b3c819ebf0db51c1b2b734922
